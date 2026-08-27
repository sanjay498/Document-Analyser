import os
from pathlib import Path
import docx

def create_sample_docx_template(output_path: str):
    doc = docx.Document()
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("LEGAL OPINION & TITLE VERIFICATION")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = docx.shared.Pt(16)
    
    doc.add_paragraph() # Spacer
    
    # Details Paragraphs
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Borrower Name: ")
    r1.bold = True
    r1_val = p1.add_run("{{BORROWER_NAME}}")
    r1_val.font.color.rgb = docx.shared.RGBColor(0, 51, 102)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Document Number: ")
    r2.bold = True
    p2.add_run("{{DOCUMENT_NUMBER}}")

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Property Description: ")
    r3.bold = True
    p3.add_run("{{PROPERTY_DESCRIPTION}}")

    p4 = doc.add_paragraph()
    r4 = p4.add_run("Date of Opinion: ")
    r4.bold = True
    p4.add_run("{{DATE}}")

    # Table
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field Name"
    hdr_cells[1].text = "Value"
    
    row_cells = table.rows[1].cells
    row_cells[0].text = "Borrower Full Reference"
    row_cells[1].text = "{{BORROWER_NAME}}"

    out_p = Path(output_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    print(f"Created sample template at {out_p}")

def create_sample_source_docx(output_path: str):
    doc = docx.Document()
    doc.add_heading("PROPERTY TITLE & REGISTRATION RECORD", level=1)
    
    doc.add_paragraph("Name of the Borrower : K.MUTHULAKSHMI, W/o G.Kumar")
    doc.add_paragraph("registered as Document No: 1277/1987")
    doc.add_paragraph("The properties in S.F.No.245/1B situated at Main Road...")
    doc.add_paragraph("Date of Execution : 05.05.1987")

    out_p = Path(output_path)
    out_p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_p))
    print(f"Created sample source doc at {out_p}")

if __name__ == "__main__":
    fixtures_dir = Path(__file__).parent / "fixtures"
    create_sample_docx_template(str(fixtures_dir / "sample_legal_opinion.docx"))
    create_sample_source_docx(str(fixtures_dir / "sample_source_record.docx"))
