import os
import re
from pathlib import Path
from typing import Dict, List, Tuple
import docx
from docx.shared import Pt, RGBColor
from app.utils.logger import logger

class DocxTemplateEngine:
    """
    DOCX generation engine that replaces dynamic fields while strictly preserving
    the original document's layout, formatting, fonts, colors, line spacing, tables, headers, and footers.
    """

    @staticmethod
    def extract_placeholders(docx_path: str) -> List[Tuple[str, str]]:
        """
        Extracts dynamic placeholders like {{FIELD_NAME}} or {{ FIELD_NAME }} from DOCX.
        Returns a list of tuples: (field_name, placeholder_str)
        """
        doc = docx.Document(docx_path)
        placeholders = set()
        pattern = re.compile(r'\{\{\s*([A-Za-z0-9_\-]+)\s*\}\}')

        def check_text(text: str):
            for match in pattern.finditer(text):
                placeholders.add((match.group(1).upper(), match.group(0)))

        # Scan paragraphs
        for p in doc.paragraphs:
            check_text(p.text)

        # Scan tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        check_text(p.text)

        # Scan headers & footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    check_text(p.text)
            if section.footer:
                for p in section.footer.paragraphs:
                    check_text(p.text)

        return list(placeholders)

    def generate_document(
        self,
        template_docx_path: str,
        output_docx_path: str,
        field_values: Dict[str, str]
    ) -> str:
        """
        Loads template, replaces placeholders with extracted values, and saves new file.
        """
        if not Path(template_docx_path).exists():
            raise FileNotFoundError(f"Template DOCX not found: {template_docx_path}")

        doc = docx.Document(template_docx_path)

        # Process all paragraphs
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, field_values)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph(p, field_values)

        # Process headers & footers
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    self._replace_in_paragraph(p, field_values)
            if section.footer:
                for p in section.footer.paragraphs:
                    self._replace_in_paragraph(p, field_values)

        # Save generated document
        output_path = Path(output_docx_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"Generated DOCX successfully saved to {output_path}")

        return str(output_path.resolve())

    def _replace_in_paragraph(self, paragraph, field_values: Dict[str, str]):
        p_text = paragraph.text
        if not p_text or "{{" not in p_text:
            return

        # Build replacement map
        replacements = {}
        for key, val in field_values.items():
            replacement_val = str(val) if val is not None else ""
            replacements["{{" + key + "}}"] = replacement_val
            replacements["{{ " + key + " }}"] = replacement_val

        # Attempt 1: Replace inside runs if single run contains placeholder
        replaced_in_runs = False
        for run in paragraph.runs:
            for ph, sub_val in replacements.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, sub_val)
                    replaced_in_runs = True

        # Attempt 2: If placeholder spans multiple runs, replace in paragraph text and inherit first run format
        if not replaced_in_runs:
            new_text = p_text
            for ph, sub_val in replacements.items():
                if ph in new_text:
                    new_text = new_text.replace(ph, sub_val)

            if new_text != p_text:
                # Capture first run's font formatting if available
                first_run = paragraph.runs[0] if paragraph.runs else None
                font_name = first_run.font.name if first_run else None
                font_size = first_run.font.size if first_run else None
                bold = first_run.bold if first_run else None
                italic = first_run.italic if first_run else None
                font_color = first_run.font.color.rgb if first_run and first_run.font.color else None

                # Clear paragraph runs
                paragraph.text = ""
                new_run = paragraph.add_run(new_text)

                # Re-apply formatting
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                if bold is not None:
                    new_run.bold = bold
                if italic is not None:
                    new_run.italic = italic
                if font_color:
                    new_run.font.color.rgb = font_color

docx_engine = DocxTemplateEngine()
