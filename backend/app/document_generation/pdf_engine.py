import os
import subprocess
import shutil
from pathlib import Path
import docx
from app.utils.logger import logger

class PdfGenerationEngine:
    def convert_docx_to_pdf(self, docx_path: str, output_pdf_path: str) -> str:
        input_path = Path(docx_path)
        out_path = Path(output_pdf_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        # Method 1: Check LibreOffice / soffice CLI
        soffice_path = shutil.which("libreoffice") or shutil.which("soffice")
        if soffice_path:
            try:
                cmd = [
                    soffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(out_path.parent),
                    str(input_path)
                ]
                res = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                converted_file = out_path.parent / (input_path.stem + ".pdf")
                if converted_file.exists():
                    if converted_file != out_path:
                        shutil.move(str(converted_file), str(out_path))
                    logger.info(f"PDF converted successfully via LibreOffice: {out_path}")
                    return str(out_path.resolve())
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {e}")

        # Method 2: Fallback ReportLab PDF renderer
        logger.info("Using ReportLab fallback for PDF generation.")
        return self._generate_reportlab_pdf(docx_path, output_pdf_path)

    def _generate_reportlab_pdf(self, docx_path: str, output_pdf_path: str) -> str:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        out_path = Path(output_pdf_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=54, bottomMargin=54
        )

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        normal_style.fontSize = 11
        normal_style.leading = 14

        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading1'],
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#1e1b4b'),
            spaceAfter=12
        )

        story = []
        docx_doc = docx.Document(docx_path)

        for p in docx_doc.paragraphs:
            text = p.text.strip()
            if not text:
                story.append(Spacer(1, 8))
                continue

            # Check if heading
            if p.style and 'Heading' in p.style.name:
                story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))

        for t in docx_doc.tables:
            table_data = []
            for row in t.rows:
                row_data = [Paragraph(cell.text.strip(), normal_style) for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                    ('PADDING', (0,0), (-1,-1), 6),
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 12))

        doc.build(story)
        logger.info(f"Generated PDF successfully saved via ReportLab to {out_path}")
        return str(out_path.resolve())

pdf_engine = PdfGenerationEngine()
