import os
from pathlib import Path
from typing import List
import docx
from PIL import Image

from app.ocr.layout import DocumentContext, TextElement
from app.utils.logger import logger

class DocumentExtractor:
    def __init__(self):
        self.easyocr_reader = None

    def _get_easyocr_reader(self):
        if self.easyocr_reader is None:
            try:
                import easyocr
                self.easyocr_reader = easyocr.Reader(['en'], gpu=False)
            except Exception as e:
                logger.warning(f"EasyOCR initialization skipped/failed: {e}")
                self.easyocr_reader = False
        return self.easyocr_reader if self.easyocr_reader is not False else None

    def extract(self, file_path: str) -> DocumentContext:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext in ['.pdf']:
            return self._extract_pdf(path)
        elif ext in ['.docx', '.doc']:
            return self._extract_docx(path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self._extract_image(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _extract_pdf(self, path: Path) -> DocumentContext:
        elements: List[TextElement] = []
        full_text_lines = []
        page_count = 1

        # Attempt 1: pdfplumber for spatial bounding boxes
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                for page_idx, page in enumerate(pdf.pages, start=1):
                    words = page.extract_words()
                    if words:
                        for word in words:
                            elements.append(TextElement(
                                text=word['text'],
                                page_number=page_idx,
                                bbox=[float(word['x0']), float(word['top']), float(word['x1']), float(word['bottom'])],
                                confidence=0.99
                            ))
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            full_text_lines.append(page_text)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed/fallback: {e}")

        # Attempt 2: PyPDF fallback if pdfplumber extracted nothing
        if not full_text_lines:
            try:
                import pypdf
                reader = pypdf.PdfReader(str(path))
                page_count = len(reader.pages)
                for page_idx, page in enumerate(reader.pages, start=1):
                    text = page.extract_text() or ""
                    if text.strip():
                        full_text_lines.append(text)
                        for line in text.splitlines():
                            if line.strip():
                                elements.append(TextElement(
                                    text=line.strip(),
                                    page_number=page_idx,
                                    bbox=[0.0, 0.0, 0.0, 0.0],
                                    confidence=0.95
                                ))
            except Exception as e:
                logger.warning(f"pypdf extraction failed: {e}")

        full_text = "\n".join(full_text_lines)
        return DocumentContext(
            file_name=path.name,
            file_type="pdf",
            total_pages=page_count,
            full_text=full_text,
            elements=elements
        )

    def _extract_docx(self, path: Path) -> DocumentContext:
        elements: List[TextElement] = []
        full_text_lines = []

        doc = docx.Document(str(path))
        for p_idx, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                full_text_lines.append(text)
                elements.append(TextElement(
                    text=text,
                    page_number=1,
                    bbox=[0.0, float(p_idx * 20), 500.0, float(p_idx * 20 + 15)],
                    confidence=1.0
                ))

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    line = " | ".join(row_texts)
                    full_text_lines.append(line)
                    elements.append(TextElement(
                        text=line,
                        page_number=1,
                        bbox=[0.0, 0.0, 0.0, 0.0],
                        confidence=1.0
                    ))

        return DocumentContext(
            file_name=path.name,
            file_type="docx",
            total_pages=1,
            full_text="\n".join(full_text_lines),
            elements=elements
        )

    def _extract_image(self, path: Path) -> DocumentContext:
        elements: List[TextElement] = []
        full_text_lines = []

        # Attempt EasyOCR if available
        reader = self._get_easyocr_reader()
        if reader:
            try:
                results = reader.readtext(str(path))
                for bbox, text, prob in results:
                    text_str = str(text).strip()
                    if text_str:
                        # bbox format from EasyOCR: [[x0,y0], [x1,y0], [x1,y1], [x0,y1]]
                        x0 = float(bbox[0][0])
                        y0 = float(bbox[0][1])
                        x1 = float(bbox[2][0])
                        y1 = float(bbox[2][1])
                        elements.append(TextElement(
                            text=text_str,
                            page_number=1,
                            bbox=[x0, y0, x1, y1],
                            confidence=float(prob)
                        ))
                        full_text_lines.append(text_str)
            except Exception as e:
                logger.warning(f"EasyOCR failed on image {path}: {e}")

        # Fallback if no text extracted (e.g. pytesseract or basic info)
        if not full_text_lines:
            try:
                import pytesseract
                img = Image.open(path)
                text = pytesseract.image_to_string(img)
                if text.strip():
                    full_text_lines.append(text.strip())
                    for line in text.splitlines():
                        if line.strip():
                            elements.append(TextElement(
                                text=line.strip(),
                                page_number=1,
                                bbox=[0.0, 0.0, 0.0, 0.0],
                                confidence=0.85
                            ))
            except Exception:
                pass

        full_text = "\n".join(full_text_lines)
        return DocumentContext(
            file_name=path.name,
            file_type="image",
            total_pages=1,
            full_text=full_text,
            elements=elements
        )

document_extractor = DocumentExtractor()
